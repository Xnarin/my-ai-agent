from datetime import datetime
from pathlib import Path
from docx import Document

import storage
from models import Transaction

def add_transaction(type, category, amount, description, date):

    if not isinstance(amount, int) or amount <= 0:
        return {
            "ok": False,
            "error": "금액은 0보다 큰 정수여야 합니다.",
        }

    # 날짜 검사 형식 준수
    try:
        datetime.strptime(date, "%Y-%m-%d")
    except ValueError:
        return {
            "ok": False,
            "error": "날짜는 YYYY-MM-DD 형식이어야 합니다.",
        }
        
    id = f"T{len(storage.TRANSACTIONS) + 1:03}"

    storage.TRANSACTIONS[id] = Transaction(
        transaction_id=id,
        transaction_type=type,
        category=category,
        amount=amount,
        description=description,
        transaction_date=date,
    )

    storage.save_data()

    return {"ok": True,"message": "거래가 저장되었습니다.","transaction": storage.TRANSACTIONS[id].model_dump(mode="json"),}


def search_transactions(date=None, category=None, description=None):
    # 조건에 맞는 구조체 저장
    results = []

    for transaction in storage.TRANSACTIONS.values():
        if date and str(transaction.transaction_date) != date:
            continue

        if category and category not in transaction.category:
            continue

        if description and description.lower() not in transaction.description.lower():
            continue
        
        # model_dump(mode="json")는 pydantic을 json으로 변환
        results.append(transaction.model_dump(mode="json"))

    return {"ok": True, "count": len(results),"transactions": results,}


def update_transaction(id,type=None,category=None,amount=None,description=None,date=None,):
    transaction = storage.TRANSACTIONS.get(id)

    if transaction is None:
        return {"ok": False,"error": "해당 거래 ID가 없습니다.",}

    data = transaction.model_dump()

    if type is not None:
        data["transaction_type"] = type

    if category is not None:
        data["category"] = category

    if amount is not None:
        data["amount"] = amount

    if description is not None:
        data["description"] = description

    if date is not None:
        data["transaction_date"] = date

    storage.TRANSACTIONS[id] = Transaction(**data)

    storage.save_data()

    return {"ok": True, "message": "거래가 수정되었습니다.","transaction": storage.TRANSACTIONS[id].model_dump(mode="json"),}


def set_budget(month, category, amount):
    if month not in storage.BUDGETS:
        storage.BUDGETS[month] = {}

    storage.BUDGETS[month][category] = amount

    storage.save_data()

    return {"ok": True,"month": month,"category": category,"amount": amount,}


def check_budget(month, category):
    budget = storage.BUDGETS.get(month, {}).get(category)

    if budget is None:
        return {
            "ok": False,
            "error": "설정된 예산이 없습니다.",
        }

    spent = 0

    for transaction in storage.TRANSACTIONS.values():
        is_same_month = str(transaction.transaction_date).startswith(month)
        is_same_category = transaction.category == category
        is_expense = transaction.transaction_type == "지출"

        if is_same_month and is_same_category and is_expense:
            spent += transaction.amount

    return {
        "ok": True,
        "month": month,
        "category": category,
        "budget": budget,
        "spent": spent,
        "remaining": budget - spent,
    }


def export_monthly_doc(month):
    doc = Document()
    doc.add_heading(f"{month} 거래 내역", level=1)

    count = 0

    for transaction in storage.TRANSACTIONS.values():
        if str(transaction.transaction_date).startswith(month):
            text = (
                f"{transaction.transaction_date} | "
                f"{transaction.transaction_type} | "
                f"{transaction.category} | "
                f"{transaction.amount:,}원 | "
                f"{transaction.description}"
            )

            doc.add_paragraph(text)
            count += 1

    if count == 0:
        return {
            "ok": False,
            "error": f"{month} 거래 내역이 없습니다.",
        }

    doc.add_heading("카테고리별 예산 현황", level=2)

    for category, budget in storage.BUDGETS.get(month, {}).items():
        spent = 0

        for transaction in storage.TRANSACTIONS.values():
            if (
                str(transaction.transaction_date).startswith(month)
                and transaction.transaction_type == "지출"
                and transaction.category == category
            ):
                spent += transaction.amount

        doc.add_paragraph(
            f"{category} | "
            f"예산: {budget:,}원 | "
            f"지출: {spent:,}원 | "
            f"남은 금액: {budget - spent:,}원"
        )

    output_dir = Path("reports")
    output_dir.mkdir(exist_ok=True)

    output_path = output_dir / f"{month}_거래내역.docx"
    doc.save(output_path)

    return {
        "ok": True,
        "message": "문서가 저장되었습니다.",
        "file_path": str(output_path.resolve()),
    }